import os
import pandas as pd
from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from num2words import num2words
from datetime import datetime
import getpass

# Document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


class gerador_subsidio:

    def __init__(self):

        self.window = Tk()
        self.window.title('Gerador de Subsídio')
        self.window.minsize(width=450, height=200)
        self.window.resizable(False, False)
        self.window.config(bg='#020e25')

        self.local_i = None
        self.doc_subsidio = None

        self.cabecalho = PhotoImage(file='imagens/cabeça.png')
        self.label1 = Label(image=self.cabecalho, bg='#020e25')
        self.label1.pack()

        self.label3 = Label(bg='#020e25')
        self.label3.pack(fill='x')

        self.t_instrumentos = Label(self.label3, text='Planilha de Instrumentos:', bg='#020e25', fg='white',
                                    pady=5)
        self.t_instrumentos.pack(side='left')

        self.b_link_instrumentos = Button(self.label3, text='Link', bg='gray', fg='white',
                                          command=lambda: self.open_link(1))
        self.b_link_instrumentos.pack(side='left', padx=3)

        self.b_instrumentos = Button(self.label3, command=self.directory_excel_i, text='Selecione o arquivo '
                                                                                       'da planilha')
        self.b_instrumentos.pack(side='left', ipadx=10)

        self.b_load_instrumentos = Button(self.label3, command=self.load_data_instrumentos, text='Carregar dados',
                                          bg='#444444', fg='white')
        self.b_load_instrumentos.pack(side='left', padx=10)

        self.label4 = Label(bg='#020e25')
        self.label4.pack(fill='x')

        self.cmb = Label(self.label4, text='Planilha de Casas (CMB): ', bg='#020e25', fg='white', pady=5)
        self.cmb.pack(side='left')

        self.b_link_cmb = Button(self.label4, text='Link', bg='gray', fg='white', command=lambda: self.open_link(0))
        self.b_link_cmb.pack(side='left', padx=3)

        self.b_cmb = Button(self.label4, command=self.directory_excel_c, text='Selecione o arquivo da planilha')
        self.b_cmb.pack(side='left', ipadx=10)

        self.b_load_cmb = Button(self.label4, command=self.load_data_cmb, text='Carregar dados', bg='#444444',
                                 fg='white')
        self.b_load_cmb.pack(side='left', padx=10)

        # Label 5

        self.label_selecao = Label(bg='#020e25', text='Seleção das UFs', fg='white', pady=10)
        self.label_selecao.pack()

        self.label_norte = Label(bg='#020e25', pady=10)
        self.label_norte.pack(fill='x')

        self.label_nordeste = Label(bg='#020e25', pady=10)
        self.label_nordeste.pack(fill='x')

        self.label_centro = Label(bg='#020e25', pady=10)
        self.label_centro.pack(fill='x')

        self.label_sudeste = Label(bg='#020e25', pady=10)
        self.label_sudeste.pack(fill='x')

        self.label_sul = Label(bg='#020e25', pady=10)
        self.label_sul.pack(fill='x')

        self.estados = {'AC': 'Acre', 'AL': 'Alagoas', 'AP': 'Amapá', 'AM': 'Amazonas', 'BA': 'Bahia', 'CE': 'Ceará',
                        'DF': 'Distrito Federal', 'ES': 'Espírito Santo', 'GO': 'Goiás', 'MA': 'Maranhão',
                        'MT': 'Mato Grosso', 'MS': 'Mato Grosso do Sul', 'MG': 'Minas Gerais', 'PA': 'Pará',
                        'PB': 'Paraíba',
                        'PR': 'Paraná', 'PE': 'Pernambuco', 'PI': 'Piauí', 'RJ': 'Rio de Janeiro',
                        'RN': 'Rio Grande do Norte', 'RS': 'Rio Grande do Sul', 'RO': 'Rondônia', 'RR': 'Roraima',
                        'SC': 'Santa Catarina', 'SP': 'São Paulo', 'SE': 'Sergipe', 'TO': 'Tocantins'}
        self.month_name = {'1': 'janeiro', '2': 'fevereiro', '3': 'março', '4': 'abril', '5': 'maio', '6': 'junho',
                           '7': 'julho', '8': 'agosto', '9': 'setembro', '10': 'outubro', '11': 'novembro',
                           '12': 'dezembro'}

        self.label_regiao_norte = Label(self.label_norte, text='Região Norte:', bg='#020e25', fg='white')
        self.label_regiao_norte.pack(side='left')

        self.label_regiao_nordeste = Label(self.label_nordeste, text='Região Nordeste:', bg='#020e25', fg='white')
        self.label_regiao_nordeste.pack(side='left')

        self.label_regiao_centro_oeste = Label(self.label_centro, text='Região Centro-Oeste:', bg='#020e25', fg='white')
        self.label_regiao_centro_oeste.pack(side='left')

        self.label_regiao_sudeste = Label(self.label_sudeste, text='Região Sudeste:', bg='#020e25', fg='white')
        self.label_regiao_sudeste.pack(side='left')

        self.label_regiao_sul = Label(self.label_sul, text='Região Sul:', bg='#020e25', fg='white')
        self.label_regiao_sul.pack(side='left')

        self.v_ac = IntVar()
        self.b_ac = Checkbutton(self.label_norte, text='AC', variable=self.v_ac, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ac.pack(side='left')

        self.v_al = IntVar()
        self.b_al = Checkbutton(self.label_nordeste, text='AL', variable=self.v_al, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_al.pack(side='left')

        self.v_ap = IntVar()
        self.b_ap = Checkbutton(self.label_norte, text='AP', variable=self.v_ap, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ap.pack(side='left')

        self.v_am = IntVar()
        self.b_am = Checkbutton(self.label_norte, text='AM', variable=self.v_am, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_am.pack(side='left')

        self.v_ba = IntVar()
        self.b_ba = Checkbutton(self.label_nordeste, text='BA', variable=self.v_ba, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ba.pack(side='left')

        self.v_ce = IntVar()
        self.b_ce = Checkbutton(self.label_nordeste, text='CE', variable=self.v_ce, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ce.pack(side='left')

        self.v_df = IntVar(value=1)
        self.b_df = Checkbutton(self.label_centro, text='DF', variable=self.v_df, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_df.pack(side='left')

        self.v_es = IntVar()
        self.b_es = Checkbutton(self.label_sudeste, text='ES', variable=self.v_es, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_es.pack(side='left')

        self.v_go = IntVar()
        self.b_go = Checkbutton(self.label_centro, text='GO', variable=self.v_go, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_go.pack(side='left')

        self.v_ma = IntVar()
        self.b_ma = Checkbutton(self.label_nordeste, text='MA', variable=self.v_ma, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ma.pack(side='left')

        self.v_ms = IntVar()
        self.b_ms = Checkbutton(self.label_centro, text='MS', variable=self.v_ms, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ms.pack(side='left')

        self.v_mt = IntVar()
        self.b_mt = Checkbutton(self.label_centro, text='MT', variable=self.v_mt, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_mt.pack(side='left')

        self.v_mg = IntVar()
        self.b_mg = Checkbutton(self.label_sudeste, text='MG', variable=self.v_mg, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_mg.pack(side='left')

        self.v_pa = IntVar()
        self.b_pa = Checkbutton(self.label_norte, text='PA', variable=self.v_pa, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_pa.pack(side='left')

        self.v_pb = IntVar()
        self.b_pb = Checkbutton(self.label_nordeste, text='PB', variable=self.v_pb, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_pb.pack(side='left')

        self.v_pr = IntVar()
        self.b_pr = Checkbutton(self.label_sul, text='PR', variable=self.v_pr, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_pr.pack(side='left')

        self.v_pe = IntVar()
        self.b_pe = Checkbutton(self.label_nordeste, text='PE', variable=self.v_pe, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_pe.pack(side='left')

        self.v_pi = IntVar()
        self.b_pi = Checkbutton(self.label_nordeste, text='PI', variable=self.v_pi, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_pi.pack(side='left')

        self.v_rj = IntVar()
        self.b_rj = Checkbutton(self.label_sudeste, text='RJ', variable=self.v_rj, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_rj.pack(side='left')

        self.v_rn = IntVar()
        self.b_rn = Checkbutton(self.label_nordeste, text='RN', variable=self.v_rn, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_rn.pack(side='left')

        self.v_rs = IntVar()
        self.b_rs = Checkbutton(self.label_sul, text='RS', variable=self.v_rs, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_rs.pack(side='left')

        self.v_ro = IntVar()
        self.b_ro = Checkbutton(self.label_norte, text='RO', variable=self.v_ro, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_ro.pack(side='left')

        self.v_rr = IntVar()
        self.b_rr = Checkbutton(self.label_norte, text='RR', variable=self.v_rr, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_rr.pack(side='left')

        self.v_sc = IntVar()
        self.b_sc = Checkbutton(self.label_sul, text='SC', variable=self.v_sc, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_sc.pack(side='left')

        self.v_sp = IntVar()
        self.b_sp = Checkbutton(self.label_sudeste, text='SP', variable=self.v_sp, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_sp.pack(side='left')

        self.v_se = IntVar()
        self.b_se = Checkbutton(self.label_nordeste, text='SE', variable=self.v_se, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_se.pack(side='left')

        self.v_to = IntVar()
        self.b_to = Checkbutton(self.label_norte, text='TO', variable=self.v_to, bg='#020e25', fg='white',
                                activebackground='black', activeforeground='white', selectcolor='black')
        self.b_to.pack(side='left')

        self.lista_variavel = [self.v_ac, self.v_al, self.v_ap, self.v_am, self.v_ba, self.v_ce, self.v_df, self.v_es,
                               self.v_go, self.v_ma, self.v_ms, self.v_mt, self.v_mg, self.v_pa, self.v_pb, self.v_pr,
                               self.v_pe, self.v_pi, self.v_rj, self.v_rn, self.v_rs, self.v_ro, self.v_rr, self.v_sc,
                               self.v_sp, self.v_se, self.v_to]

        # Limpar seleções
        self.label_limpar_selecao = Label(bg='#020e25')
        self.label_limpar_selecao.pack(fill='x')

        self.b_load_cmb = Button(self.label_sudeste, text='Limpar Seleções', bg='#444444', fg='white',
                                 command=self.clear_selections)
        self.b_load_cmb.pack(padx=32, side='right')

        # Diretório para Salvamento
        self.label6 = Label(bg='#020e25', pady=5)
        self.label6.pack(fill='x')

        self.t_diretorio = Label(self.label6, text='Pasta para salvamento: ', bg='#020e25', fg='white')
        self.t_diretorio.pack(side='left', pady=10)

        self.local_save = 'C:/Users'

        self.b_save = Button(self.label6, text='Escolha a pasta para o salvamento dos arquivos',
                             command=self.directory_save)
        self.b_save.pack(side='left')

        self.label_cabecalho = Label(bg='#020e25', pady=5)
        self.label_cabecalho.pack(fill='x')

        self.t_cabecalho = Label(self.label_cabecalho, text='Cabeçalho:', bg='#020e25', fg='white')
        self.t_cabecalho.pack(side='left')

        self.v_cabecalho = IntVar(value=1)
        self.b_cabecalho = Checkbutton(self.label_cabecalho, bg='#020e25', fg='white', variable=self.v_cabecalho,
                                       activebackground='black', activeforeground='white', selectcolor='black').pack(
                                        side='left')

        self.label7 = Label(bg='#020e25')
        self.label7.pack(fill='x')

        self.b_manual = Button(self.label7, text='Abrir manual', bg='#444444', fg='white',
                               command=lambda: self.open_link(2))
        self.b_manual.pack(side='right', pady=2, ipadx=2, ipady=2, padx=35)

        self.b_gerar = Button(self.label7, text='GERAR SUBSÍDIO', command=self.gerar_subsidio, bg='#23abb2',
                              fg='white')
        self.b_gerar.pack(side='right', pady=10, ipadx=10, ipady=5, padx=30)

        self.image_logo = PhotoImage(file='imagens/logo dev_snpm_mmfdh.png')
        self.label_logo = Label(bg='#020e25', image=self.image_logo)
        self.label_logo.pack(fill='x')

        self.window.mainloop()

    def open_explorer(self):
        os.system(f'start {self.local_save}')

    @staticmethod
    def open_link(op):
        import webbrowser
        if op == 1:
            link = 'google.com.br'
        elif op == 2:
            link = 'google.com.br'
        else:
            link = 'google.com.br'
        webbrowser.open(link)
        return None

    def directory_excel_i(self):
        filetypes = (('Planilha do Microsoft Excel', '*.xlsx'), ('All files', '*.*'))
        self.local_i = filedialog.askopenfilename(filetypes=filetypes, initialdir='C:/Users/renan.santos/Downloads')
        print(self.local_i)
        return None

    def directory_excel_c(self):
        filetypes = (('Planilha do Microsoft Excel', '*.xlsx'), ('All files', '*.*'))
        self.local_c = filedialog.askopenfilename(filetypes=filetypes, initialdir='C:/Users/renan.santos/Downloads')
        print(self.local_c)
        return None

    def number_to_long_number(self, number_p):
        if number_p.find('.') != -1:
            number_p = number_p.split('.')
            number_p1 = int(number_p[0].replace('.', ''))
            number_p2 = int(number_p[1])
        else:
            number_p1 = int(number_p.replace('.', ''))
            number_p2 = 0
        if number_p1 == 1:
            aux1 = ' real'
        else:
            aux1 = ' reais'
        if number_p2 == 1:
            aux2 = ' centavo'
        else:
            aux2 = ' centavos'
        if number_p1 > 0:
            text1 = num2words(number_p1, lang='pt_BR') + str(aux1)
        else:
            text1 = ''
        if number_p2 > 0:
            text2 = num2words(number_p2, lang='pt_BR') + str(aux2)
        else:
            text2 = ''
        if (number_p1 > 0) and (number_p2 > 0):
            result = text1 + ' e ' + text2
        else:
            result = text1 + text2
        return result

    def directory_save(self):
        user_name = getpass.getuser()
        self.local_save = filedialog.askdirectory(initialdir=f'C:/Users/{user_name}/Documents')
        print(self.local_save)
        return None

    def gerar_subsidio(self):
        lista_uf = self.get_uf_selected()
        if len(lista_uf) > 0:
            for uf in lista_uf:
                self.gerar_doc(uf)
        else:
            self.error_window('sds')

    @staticmethod
    def error_window(text):
        w = Toplevel()
        w.title('ERRO!')
        w.geometry('300x300+300+300')
        w.resizable(False, False)
        w.config(bg='#444444')

        text = ttk.Label(w, text=str(text), background='#444444', foreground='white')
        text.pack()

        # close_button = ttk.Button(w, text='Ok', command=w.destroy())
        # close_button.pack()

    def load_data_instrumentos(self):
        self.data_convenios = pd.read_excel(f'{self.local_i}', sheet_name='Convênios 2018.2022').loc[
                              :105, :]
        self.data_act = pd.read_excel(f'{self.local_i}', sheet_name='ACT 2018-2022')
        self.data_contratos = pd.read_excel(f'{self.local_i}',
                                            sheet_name='Contratos de Repasse.2018.2022').loc[
                              :40, :]
        self.data_teds = pd.read_excel(f'{self.local_i}', sheet_name='TEDS 2018.2022').loc[:5, :]
        return None

    def load_data_cmb(self):
        self.data_candamento = pd.read_excel(self.local_c, header=1)
        data_temp1 = self.data_candamento.loc[1:3, :]
        data_temp2 = self.data_candamento.loc[5:6, :]
        data_temp3 = self.data_candamento.loc[8:25, :]
        data_temp4 = self.data_candamento.loc[27:34, :]
        data_temp5 = self.data_candamento.loc[38:39, :]
        self.data_c_andamento = pd.concat([data_temp1, data_temp2, data_temp3, data_temp4, data_temp5], axis=0)
        self.data_c_andamento['CMB/CRAM'] = self.data_c_andamento['Tipologia'].apply(lambda x: 'CRAM' if x == "_"
        else "CMB")
        return None

    def get_uf_selected(self):
        estados_min = ['ac', 'al', 'ap', 'am', 'ba', 'ce', 'df', 'es', 'go', 'ma', 'ms', 'mt', 'mg', 'pa', 'pb', 'pr',
                       'pe', 'pi', 'rj', 'rn', 'rs', 'ro', 'rr', 'sc', 'sp', 'se', 'to']
        l_uf_selecionada = []

        for num, vari in enumerate(self.lista_variavel):
            if vari.get() == 1:
                l_uf_selecionada.append(estados_min[num].upper())
        print(l_uf_selecionada)
        return l_uf_selecionada

    def clear_selections(self):
        for variavel in self.lista_variavel:
            variavel.set(value=0)
        return None

    def cabecalho1(self, uf):
        self.doc_subsidio.add_picture(r'imagens\governo_federal_logo.png', width=Inches(0.8), height=Inches(0.8))
        last_paragraph = self.doc_subsidio.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cabecalho = self.doc_subsidio.add_paragraph()
        cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cabecalho.add_run(
            f'MINISTÉRIO DA MULHER, DA FAMÍLIA E DOS DIREITOS HUMANOS (MMFDH) \nSECRETARIA NACIONAL DE POLÍTICAS PARA '
            f'MULHERES (SNPM)')
        p.bold = True
        self.doc_subsidio.add_paragraph()

    def cabecalho2(self, uf):
        data_criacao = datetime.today()

        cabecalho = self.doc_subsidio.add_paragraph()
        cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = cabecalho.add_run(
            f'{self.estados[uf].upper()}')
        p.bold = True

        introducao = self.doc_subsidio.add_paragraph(
            f'No Estado de {self.estados[uf].capitalize()}, em {data_criacao.day} de '
            f'{self.month_name[str(data_criacao.month)]} '
            f'de {data_criacao.year}, destacamos os seguintes instrumentos de transferência voluntária afetos ao '
            f'enfrentamento à violência contra as mulheres.')
        introducao.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def gerar_doc(self, uf):

        def valor_formatado(nome_t, nome_r, linha):
            vr = f'{linha[nome_r]:_.2f}'
            t_vr = vr.replace('.', ',').replace('_', '.')
            vt = f'R$ {linha[nome_t]:_.2f}'
            t_vt = vt.replace('.', ',').replace('_', '.')
            return [t_vt, t_vr]

        self.doc_subsidio = Document()
        style = self.doc_subsidio.styles['Normal']
        f = style.font

        f.name = 'Calibri'
        f.size = Pt(12)

        # Base de dados
        data_temp_conv = self.data_convenios.loc[self.data_convenios['UF'] == uf, :]
        data_temp_ct = self.data_contratos.loc[self.data_contratos['UF'] == uf, :]
        data_temp_ted = self.data_teds.loc[self.data_teds['UF'] == uf, :]
        data_temp_act = self.data_act.loc[self.data_act['UF'] == uf, :]
        data_temp_cb = self.data_c_andamento.loc[(self.data_c_andamento['Estado '] == uf) &
                                                 (self.data_c_andamento['CMB/CRAM'] == 'CMB'), :]

        dic_cmbs, lista_nsincov = {}, []
        if data_temp_cb.shape[0] > 0:
            for c, _ in enumerate(data_temp_cb.values):
                linha = data_temp_cb.iloc[c]
                lista_nsincov.append(linha["Siconv"].strip())
                dic_cmbs[f'{linha["Siconv"].strip()}'] = [linha["EM CONSTRUÇÃO: SIM ou NÃO"],
                                                          linha["Previsão de Entrega (obra concluída)"],
                                                          linha["Endereço"], linha["Contato"]]

        # Cabeçalho
        if self.v_cabecalho.get() == 0:
            self.cabecalho2(uf)
        else:
            self.cabecalho1(uf)
            self.cabecalho2(uf)

        # Convênios
        if data_temp_conv.shape[0] > 0:
            for c, _ in enumerate(data_temp_conv.values):
                linha = data_temp_conv.iloc[c]
                numero = linha['Nº PLATAFORMA +BRASIL']
                lista_valores = valor_formatado(nome_t="VALOR DO CONVÊNIO", nome_r="VALOR DE REPASSE", linha=linha)
                vt_conv_text = self.number_to_long_number(str(linha["VALOR DO CONVÊNIO"]))

                nome = 'Convênio'
                if linha['INSTRUMENTO'].upper() == 'TERMO DE FOMENTO':
                    nome = 'Termo de Fomento'

                convenios_texto = self.doc_subsidio.add_paragraph(style='List Number')
                convenio_numero = convenios_texto.add_run(f'{nome} nº {numero}\n')
                convenio_numero.bold = True
                convenio_numero.underline = True
                if str(linha['Município']) == 'nan':
                    convenio_municipio = convenios_texto.add_run('Abrangência: ')
                    convenio_municipio.bold = True
                    convenios_texto.add_run('Estadual.\n')
                else:
                    convenio_municipio = convenios_texto.add_run(f'Município(s): ')
                    convenio_municipio.bold = True
                    convenios_texto.add_run(f'{linha["Município"]}.\n')

                convenio_l = convenios_texto.add_run(f'Objeto: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{linha["OBJETO"]}')
                if linha['OBJETO'][-1] != '.':
                    convenios_texto.add_run('.')

                convenio_l = convenios_texto.add_run(f'\nConvenente: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{linha["CONVENENTE"]}')

                # convenio_l = convenios_texto.add_run(f'\nPúblico-alvo: ')
                # convenio_l.bold = True
                # convenios_texto.add_run(f'{linha["PÚBLICO-ALVO/MULHERES BENEFICIADAS"]}')
                #
                # convenio_l = convenios_texto.add_run(f'\nBeneficiários (ano): ')
                # convenio_l.bold = True
                # convenios_texto.add_run(f'{linha["QUANTIDADE DE BENEFICIÁRIOS (ANO)"]}')

                convenio_l = convenios_texto.add_run('\nOrigem do recurso: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{linha["ORIGEM DO RECURSO"]}, ')
                emenda = str(linha['EMENDA']).split('\n')
                emenda = ' - '.join(emenda)
                convenios_texto.add_run(f'{emenda}.')

                convenio_l = convenios_texto.add_run('\nValor total: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{lista_valores[0]} ({vt_conv_text}).')

                convenio_l = convenios_texto.add_run('\nValor de repasse: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{lista_valores[1]}.')

                # Conseguir os dados da CMB do convênio

                convenio_l = convenios_texto.add_run('\nStatus: ')
                convenio_l.bold = True
                convenios_texto.add_run(f'{linha["STATUS / FASE"].strip()}.')

                self.doc_subsidio.add_paragraph()

        # Contratos de Repasse
        if data_temp_ct.shape[0] > 0:
            for c, _ in enumerate(data_temp_ct.values):
                linha = data_temp_ct.iloc[c]
                numero = linha['CONVÊNIO SICONV']
                lista_valores = valor_formatado(nome_t="VALOR TOTAL DO CONVÊNIO", nome_r="VALOR DE REPASSE",
                                                linha=linha)
                vt_cont_text = self.number_to_long_number(str(linha["VALOR TOTAL DO CONVÊNIO"]))

                contratos_texto = self.doc_subsidio.add_paragraph(style='List Number')
                contrato_numero = contratos_texto.add_run(f'Contrato de Repasse nº {numero.strip()}\n')
                contrato_numero.bold = True
                contrato_numero.underline = True
                contrato_municipio = contratos_texto.add_run(f'Município(s): ')
                contrato_municipio.bold = True
                contratos_texto.add_run(f'{linha["Município"]}.\n')
                contrato_l = contratos_texto.add_run(f'Objeto: ')
                contrato_l.bold = True
                contratos_texto.add_run(f'{linha["OBJETO"]}')
                if linha['OBJETO'][-1] != '.':
                    contratos_texto.add_run('.')

                contrato_l = contratos_texto.add_run('\nOrigem do recurso: ')
                contrato_l.bold = True
                contratos_texto.add_run(f'{linha["ORIGEM DO RECURSO"]}, ')
                emenda = str(linha['AUTOR EMENDA (Partido/UF)']).split('\n')
                emenda = ' - '.join(emenda)
                contratos_texto.add_run(f'{emenda}.')

                contrato_l = contratos_texto.add_run('\nValor total: ')
                contrato_l.bold = True
                contratos_texto.add_run(f'R$ {lista_valores[0]} ({vt_cont_text}).')

                contrato_l = contratos_texto.add_run('\nValor de repasse: ')
                contrato_l.bold = True
                contratos_texto.add_run(f'R$ {lista_valores[1]}.')

                if numero in lista_nsincov:
                    info = dic_cmbs[numero]

                    contrato_l = contratos_texto.add_run('\nEm construção? ')
                    contrato_l.bold = True

                    contratos_texto.add_run(f'{info[0].strip().capitalize()}.')

                    contrato_l = contratos_texto.add_run('\nPrevisão da obra concluída: ')
                    contrato_l.bold = True

                    if str(info[1]) == 'NaT':
                        contratos_texto.add_run('Sem previsão.')
                    else:
                        contratos_texto.add_run(f'{info[1]}.')

                    contrato_l = contratos_texto.add_run('\nEndereço: ')
                    contrato_l.bold = True
                    contratos_texto.add_run(f'{info[2].strip().capitalize()}.')

                    contrato_l = contratos_texto.add_run('\nContato: ')
                    contrato_l.bold = True
                    contratos_texto.add_run(f'{info[3].strip().capitalize()}.')

                contrato_l = contratos_texto.add_run('\nStatus: ')
                contrato_l.bold = True
                contratos_texto.add_run(f'{linha["SITUAÇÃO"].strip()}.')

                self.doc_subsidio.add_paragraph()

        # TEDs
        if data_temp_ted.shape[0] > 0:
            for c, _ in enumerate(data_temp_ted.values):
                linha = data_temp_ted.iloc[c]
                lista_valores = valor_formatado(nome_t="VALOR TOTAL DO CONVÊNIO", nome_r="VALOR DE REPASSE",
                                                linha=linha)
                vt_ted_text = self.number_to_long_number(str(linha["VALOR TOTAL DO CONVÊNIO"]))

                ted_texto = self.doc_subsidio.add_paragraph(style='List Number')
                ted_classe = ted_texto.add_run(f'TED n° {linha["PROPOSTA SICONV"]}')
                ted_classe.bold = True
                ted_classe.underline = True

                ted_negrito = ted_texto.add_run(f'\nObjeto: ')
                ted_negrito.bold = True
                ted_texto.add_run(f'{linha["OBJETO"]}')

                ted_negrito = ted_texto.add_run(f'\nValor total: ')
                ted_negrito.bold = True
                ted_texto.add_run(f'R$ {lista_valores[0]} ({vt_ted_text})')

                ted_negrito = ted_texto.add_run(f'\nSituação: ')
                ted_negrito.bold = True
                ted_texto.add_run(f'{linha["SITUAÇÃO"]}')

                self.doc_subsidio.add_paragraph()

        # ACT
        if data_temp_act.shape[0] > 0:
            for c, _ in enumerate(data_temp_act.values):
                linha = data_temp_act.iloc[c]

                act_texto = self.doc_subsidio.add_paragraph(style='List Number')
                act_classe = act_texto.add_run(f'ACT n° SEI {linha["Nº PROCESSO SEI"]}\n')
                act_classe.bold = True
                act_classe.underline = True

                act_negrito = act_texto.add_run(f'Objeto: ')
                act_negrito.bold = True
                act_texto.add_run(f'{linha["OBJETO"]}')

                self.doc_subsidio.add_paragraph()

        # Finalização
        self.doc_subsidio.add_paragraph('Permanecemos à disposição para esclarecimentos adicionais.')

        # Salvamento do arquivo
        self.doc_subsidio.save(f'{self.local_save}\{uf.upper()}_subsidio.docx')


gerador_subsidio()
